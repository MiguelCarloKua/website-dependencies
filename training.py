import sys, os, json, re
from bs4 import BeautifulSoup
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from rouge_score import rouge_scorer
from bert_score import score as bert_score
import numpy as np
import google.generativeai as genai
import requests
import csv

# --- Load environment and configure Gemini ---
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))


def fetch_court_case(url, max_retries=3):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
    }
    for attempt in range(max_retries):
        try:
            print(f"🔄 Attempt {attempt + 1}: {url}", file=sys.stderr)
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code == 200:
                return BeautifulSoup(response.content, "html.parser")
            else:
                print(f"⚠️ Status Code: {response.status_code}", file=sys.stderr)
        except Exception as e:
            print(f"❌ Request failed: {e}", file=sys.stderr)
    return None

def is_meaningful_line(line):
    line = line.strip()
    if len(line) < 5:
        return False
    if re.fullmatch(r'([xX*#.\-– ])\1{2,}', line.replace(' ', '')):
        return False
    if len(set(line.lower())) < 3:
        return False
    return True

def clean_line(line):
    line = re.sub(r'^(x\s+)+', '', line.strip(), flags=re.IGNORECASE)
    line = re.sub(r'(x\s+)+$', '', line.strip(), flags=re.IGNORECASE)
    return line.strip()

def extract_full_text(soup):
    for tag in soup.find_all(["a", "sup"]):
        tag.decompose()

    def join_paragraphs(paragraphs):
        return "\n\n".join(p.get_text(" ", strip=True) for p in paragraphs if p.get_text(strip=True))

    strategies = [
        lambda s: s.find("table", {"id": "lwphl"}).find("blockquote").find_all("p") if s.find("table", {"id": "lwphl"}) else [],
        lambda s: s.find("td", colspan="2").find_all(["p", "blockquote"]) if s.find("td", colspan="2") else [],
        lambda s: [s.find("pre")] if s.find("pre") else [],
        lambda s: s.find_all("center")[-1].find_all("p") if s.find_all("center") else [],
        lambda s: s.find("body").find_all("p") if s.find("body") else []
    ]

    for strat in strategies:
        try:
            paragraphs = strat(soup)
            text = join_paragraphs(paragraphs)
            footnote_cutoff = re.search(r'\n\s*Footnotes\s*[:\n]', text, re.IGNORECASE)
            if footnote_cutoff:
                text = text[:footnote_cutoff.start()]
            if len(text.strip()) > 300:
                return text
        except:
            continue
    return ""

def is_text_garbled(text):
    return text.count("�") > 10 or len(text.strip()) < 100


def extract_case_details(text):
    text = text.replace('\r\n', '\n')

    court_level_match = re.search(r"(SUPREME COURT|COURT OF APPEALS|SANDIGANBAYAN|REGIONAL TRIAL COURT)", text, re.IGNORECASE)
    court_level = court_level_match.group(0).strip().title() if court_level_match else "Not Found"

    division_match = re.search(r"(FIRST|SECOND|THIRD|FOURTH|FIFTH)\s+DIVISION", text, re.IGNORECASE)
    is_en_banc = "Yes" if re.search(r"\bEn\s*Banc\b", text, re.IGNORECASE) else "No"
    division = "En Banc" if is_en_banc == "Yes" else (division_match.group(0).strip().title() if division_match else "Not Found")

    gr_match = re.search(r"G\.R\. No\.\s*\d{5,}", text)
    gr_number = gr_match.group().strip() if gr_match else "Not Found"

    date_match = re.search(r"([A-Za-z]+\s\d{1,2},\s\d{4})", text)
    trial_date = date_match.group(0).strip() if date_match else "Not Found"

    lines = text.split("\n")
    party_line = next((line for line in lines if re.search(r"\bvs\.|versus\b", line, re.IGNORECASE)), "")
    party_match = re.search(r"(.+?)\s+(vs\.|versus)\s+(.+)", party_line, re.IGNORECASE)

    if party_match:
        petitioners = party_match.group(1).strip().title()
        respondents = party_match.group(3).strip().title()
    else:
        petitioners = "Not Found"
        respondents = "Not Found"

    petitioners = re.sub(r"\bManila\b|\b[A-Z]+\s+DIVISION\b|G\.R\. No\..*?\d{4}", "", petitioners, flags=re.IGNORECASE).strip(" ,")

    judge_match = re.search(r'DECISION\s*\n+([A-ZÑ ,\.\-]+?),\s*(C\.J\.|J\.)[:.]?', text, re.IGNORECASE)
    if judge_match:
        name = judge_match.group(1).strip().title()
        title = judge_match.group(2).strip()
        judge = f"{name} {title}"
    else:
        fallback_match = re.search(r'\n\s*([A-ZÑ ,\.\-]+?),\s*(C\.J\.|J\.)[:.]', text, re.IGNORECASE)
        judge = f"{fallback_match.group(1).strip().title()} {fallback_match.group(2).strip()}" if fallback_match else "Not Found"

    return {
        "G.R. Number": gr_number,
        "Date of Trial": trial_date,
        "Court Level": court_level,
        "Division": division,
        "Is En Banc": is_en_banc,
        "Petitioners": petitioners,
        "Respondents": respondents,
        "Ponente": judge
    }

def extract_case_sections_from_text(text):
    section_patterns = {
        "Facts": [
            r"(?:^|\n)\s*(?:The\s+)?("
            r"Facts|Factual Antecedents|Background|Antecedents|The Case|Statement of Facts|The Antecedent Facts|"
            r"Factual Background|Summary of Facts|Procedural History|Version of the Prosecution|"
            r"Version of the Defense|Facts and Antecedent Proceedings|Narration of Facts|"
            r"Facts of the Case|Case Background|Chronology of Events|Statement of the Case|"
            r"Case Summary|The Incident|Recital of Facts|Historical Background|"
            r"Overview of Facts"
            r")\s*(?:\n|:)"
        ],
        "Issues": [
            r"(?:^|\n)\s*(?:The\s+)?("
            r"Issue Before the Court|Issues Before the Court|Issue|Issues|Legal Issue|Questions Presented|"
            r"Statement of Issues|Points for Determination|Legal Questions|Controversy|"
            r"Questions for Resolution|Matter for Consideration|Core Issue|"
            r"Principal Issue|Pivotal Issue|Legal Questions Posed|Legal Questions Raised"
            r")\s*(?:\n|:)"
        ],
        "Ruling": [
            r"(?:^|\n)\s*(?:The\s+)?("
            r"Ruling|Decision|Held|Disposition|So Ordered|Judgment|Court['’]s Ruling|"
            r"Our Ruling|The Ruling of the Court|This Court['’]s Ruling|Ruling of the Court|"
            r"Opinion|Holding|The Ruling of this Court|Final Disposition|Resolution|"
            r"Conclusion|Finding|Adjudication|Result|Verdict|Decree"
            r")\s*(?:\n|:)"
        ]
    }

    matches = []
    for section, patterns in section_patterns.items():
        for pattern in patterns:
            for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                matches.append((match.start(), section))

    matches.sort()
    grouped = {"Facts": [], "Issues": [], "Ruling": []}

    # --- Case 1: Matching headers found ---
    if matches:
        for i, (start, section) in enumerate(matches):
            end = matches[i + 1][0] if i + 1 < len(matches) else len(text)
            section_text = text[start:end].strip()
            paragraphs = [p.strip() for p in section_text.split("\n") if p.strip()]
            cleaned = [clean_line(p) for p in paragraphs]
            meaningful = [p for p in cleaned if is_meaningful_line(p)]
            grouped[section].extend(meaningful)
        return grouped

    # --- Case 2: No headers found — fallback logic ---
    print("⚠️ No headers detected — using fallback segmentation.")
    lines = [line.strip() for line in text.split("\n") if is_meaningful_line(line)]
    total = len(lines)

    if total > 0:
        grouped["Facts"] = lines[: total // 3]
        grouped["Issues"] = lines[total // 3 : (2 * total) // 3]
        grouped["Ruling"] = lines[(2 * total) // 3 :]

    return grouped



def generate_gemini_response(prompt, text):
    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(f"{prompt}\n\n{text}")
        return response.text.strip()
    except Exception as e:
        return f"❌ Gemini error: {str(e)}"


def write_docx(summary, out_path, metadata):
    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run(f"Digest – {summary['gr_no']}")
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    table = doc.add_table(rows=0, cols=2)
    for key, value in metadata.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value
    table.style = 'Table Grid'
    doc.add_paragraph()

    for section in ["facts", "issues", "rulings"]:
        heading = doc.add_paragraph()
        run = heading.add_run(section.upper())
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(summary[section])
        doc.add_paragraph()

    doc.save(out_path)

def clean_summary_output(raw_text: str) -> str:
    """Cleans markdown, bullets, and excessive whitespace from Gemini output."""
    lines = raw_text.strip().splitlines()

    cleaned_lines = []
    for line in lines:
        line = line.strip()

        # Remove Markdown bold and italic symbols
        line = re.sub(r"[*_`]+", "", line)
        line = re.sub(r'```csv\s*', '', line)
        line = re.sub(r'```', '', line)
        line = re.sub(r'Raw Generated (Facts|Issues|Rulings)', '', line, flags=re.IGNORECASE)
        
        # Remove markdown bullets (*, -, 1., etc.)
        line = re.sub(r"^\s*[-*•]\s*", "", line)
        line = re.sub(r"^\s*\d+\.\s*", "", line)

        # Normalize colons
        line = re.sub(r"\s*:\s*", ": ", line)

        # Remove redundant spaces
        line = re.sub(r"\s{2,}", " ", line)

        if line:
            cleaned_lines.append(line)

    return "\n".join(cleaned_lines)

def compute_rouge(generated: str, reference: str):
    scorer = rouge_scorer.RougeScorer(['rouge1', 'rougeL'], use_stemmer=True)
    scores = scorer.score(reference, generated)
    return {
        "rouge-1": scores["rouge1"].fmeasure,
        "rouge-l": scores["rougeL"].fmeasure,
    }

def compute_bert_score(generated: str, reference: str):
    P, R, F1 = bert_score([generated], [reference], lang="en", rescale_with_baseline=True)
    return {
        "bert-score": F1[0].item()
    }

def evaluate_all(generated: str, reference: str):
    rouge = compute_rouge(generated, reference)
    bert = compute_bert_score(generated, reference)
    return {
        "rouge-1": round(rouge["rouge-1"], 4),
        "rouge-l": round(rouge["rouge-l"], 4),
        "bert-score": round(bert["bert-score"], 4),
    }

def main_pipeline(url: str, direction: str):
    direction = direction.lower()
    if direction not in ["forward", "backward"]:
        return { "error": "Direction must be either 'forward' or 'backward'" }

    soup = fetch_court_case(url)
    if not soup:
        return { "error": "Failed to fetch or parse the URL" }

    text = extract_full_text(soup)
    if is_text_garbled(text):
        return { "error": "Extracted text is too short or garbled" }

    base_dir = Path(__file__).resolve().parent
    with open(base_dir / "data/config.json", "r") as f:
        config = json.load(f)

    metadata = extract_case_details(text)
    sections = extract_case_sections_from_text(text)

    if direction == "backward":
        facts_input = "\n".join(sections["Ruling"]) + "\n" + "\n".join(sections["Issues"])
        issues_input = "\n".join(sections["Ruling"]) + "\n" + "\n".join(sections["Facts"])
        rulings_input = "\n".join(sections["Facts"]) + "\n" + "\n".join(sections["Issues"])
    else:
        facts_input = "\n".join(sections["Facts"])
        issues_input = "\n".join(sections["Issues"])
        rulings_input = "\n".join(sections["Ruling"])

    summary = {
        "gr_no": metadata["G.R. Number"],
        "facts": clean_summary_output(generate_gemini_response(config["FACTS"][direction.upper()]["Instructor_Extractive"], facts_input)),
        "issues": clean_summary_output(generate_gemini_response(config["ISSUES"][direction.upper()]["Judge_Extractive"], issues_input)),
        "rulings": clean_summary_output(generate_gemini_response(config["RULINGS"][direction.upper()]["Instructor_ChainOfThought"], rulings_input)),
    }

    scores = {
        "facts": evaluate_all(summary["facts"], "\n".join(sections["Facts"])),
        "issues": evaluate_all(summary["issues"], "\n".join(sections["Issues"])),
        "rulings": evaluate_all(summary["rulings"], "\n".join(sections["Ruling"])),
    }

    return {
        "summary": summary,
        "original": {
            "facts": "\n".join(sections["Facts"]),
            "issues": "\n".join(sections["Issues"]),
            "rulings": "\n".join(sections["Ruling"]),
        },
        "scores": scores,
        "metadata": metadata
    }
